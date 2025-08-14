# recruiter_matcher_dynamic.py
import streamlit as st
import re, io, json, math
import fitz  # PyMuPDF
import docx
import pandas as pd

# --------- Optional LLM (Gemini) for STRUCTURED extraction ---------
import google.generativeai as genai

# ===================== Page =====================
st.set_page_config(page_title="Recruiter Matcher (Gemini | Dynamic)", layout="centered")
st.title("📌 Terrabit Consulting – Recruiter Matcher")
st.write("Upload **one Job Description** and **multiple Resumes**. Get match scores, reasoning, and follow-ups.")

# ===================== Configure Gemini =====================
if "GEMINI_API_KEY" not in st.secrets:
    st.error("Please set GEMINI_API_KEY in Streamlit secrets.")
    st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

# JSON extractor (forces JSON so we never scrape Markdown)
gemini_json = genai.GenerativeModel(
    "gemini-1.5-flash",
    generation_config={
        "temperature": 0,
        "top_p": 0.1,
        "top_k": 1,
        "max_output_tokens": 2048,
        "response_mime_type": "application/json",
    },
)

# Short-text generator (follow-ups only; not used in scoring)
gemini_text = genai.GenerativeModel(
    "gemini-1.5-flash",
    generation_config={"temperature": 0, "top_p": 0.1, "top_k": 1, "max_output_tokens": 1024},
)

# ===================== File Readers =====================
def read_pdf(file):
    text = ""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        for p in doc:
            text += p.get_text("text")
    return text

def read_docx(file):
    d = docx.Document(file)
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    try:
        section = d.sections[0]
        for para in section.footer.paragraphs:
            parts.append(para.text)
    except Exception:
        pass
    return "\n".join(parts)

def read_file(file):
    if file.type == "application/pdf":
        return read_pdf(file)
    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return read_docx(file)
    else:
        return file.read().decode("utf-8", errors="ignore")

def normalize_text(t: str, max_chars=60000) -> str:
    if not t: return ""
    t = t.replace("\x00", " ")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()[:max_chars]

# ===================== Light NLP Helpers (no external libs) =====================
STOP = set("""
a an the and or of to in for on with without by from as at into over under be is are was were been being have has had do does did
this that these those it its their his her your our my we they he she i you them him her
about across after against among around before behind below beneath beside between beyond during except inside near outside since through
than till until up upon within within without while system systems software hardware tool tools framework frameworks language languages
""".split())

def tokenize(text):
    # lower, keep words, numbers & plus/dot
    toks = re.findall(r"[a-zA-Z0-9\+\#\.\-]+", text.lower())
    return [t for t in toks if t not in STOP and len(t) > 1]

def bow(text):
    counts = {}
    for t in tokenize(text):
        counts[t] = counts.get(t, 0) + 1
    return counts

def cosine_sim(a_counts, b_counts):
    if not a_counts or not b_counts: return 0.0
    # dot
    dot = 0.0
    for k, v in a_counts.items():
        if k in b_counts:
            dot += v * b_counts[k]
    # norms
    na = math.sqrt(sum(v*v for v in a_counts.values()))
    nb = math.sqrt(sum(v*v for v in b_counts.values()))
    if na == 0 or nb == 0: return 0.0
    return float(dot / (na * nb))

def extract_email(text):
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    return m.group() if m else "Not found"

def extract_candidate_name(text, filename):
    m = re.search(r"(?i)Candidate Name\s*[:–-]\s*(.+)", text)
    if m:
        nm = re.sub(r"[^A-Za-z \-']", " ", m.group(1)).strip()
        if 2 <= len(nm.split()) <= 4:
            return nm.title()
    m = re.search(r"(?i)Resume of\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})", text)
    if m:
        return m.group(1).strip().title()
    first = "\n".join(text.splitlines()[:15])
    m = re.search(r"(?m)^\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*$", first)
    if m:
        return m.group(1).strip().title()
    base = re.sub(r"\.(pdf|docx|txt)$", "", filename, flags=re.I)
    base = re.sub(r"[_\-]+", " ", base)
    cand = re.sub(r"[^A-Za-z \-']", " ", base).strip()
    if 2 <= len(cand.split()) <= 4:
        return cand.title()
    return "Name Not Found"

def extract_min_years(text):
    # returns the largest years requirement mentioned
    nums = re.findall(r"(?i)(?:at\s+least|min(?:imum)?)\s*(\d+)\s*\+?\s*(?:years|yrs)|(\d+)\s*\+?\s*(?:years|yrs)\s+(?:of|in)", text)
    vals = []
    for a, b in nums:
        if a: vals.append(int(a))
        elif b: vals.append(int(b))
    return max(vals) if vals else 0

def extract_years_from_resume(text):
    nums = re.findall(r"(?i)(\d+)\s*\+?\s*(?:years|yrs)", text)
    return max([int(n) for n in nums], default=0)

# ===================== LLM JSON Extraction =====================
JD_SCHEMA = {
    "role_titles": ["e.g., QA Engineer, Backend Developer, Data Engineer"],
    "required_skills": ["list of concrete hard skills (tools, frameworks, languages)"],
    "optional_skills": ["nice-to-have skills"],
    "domain_keywords": ["e.g., fintech, insurance, retail"],
    "location_keywords": ["cities/regions or 'remote'/'hybrid'"],
}

RESUME_SCHEMA = {
    "role_titles": ["candidate titles"],
    "skills": ["candidate skills (tools, frameworks, languages)"],
    "domain_keywords": ["domains mentioned"],
    "location_keywords": ["locations mentioned"],
}

def call_json(prompt: str, schema_hint: dict) -> dict:
    try:
        resp = gemini_json.generate_content(
            f"Return valid JSON only. Follow this minimal shape (keys only, types implied):\n"
            f"{json.dumps(schema_hint, indent=2)}\n\n---\n{prompt}"
        )
        txt = getattr(resp, "text", "") or "{}"
        return json.loads(txt)
    except Exception as e:
        st.error(f"Gemini JSON error: {e}")
        return {}

def to_list(x):
    if x is None: return []
    if isinstance(x, list): return [str(i).strip() for i in x if str(i).strip()]
    parts = re.split(r"[;,|\n]+", str(x))
    return [p.strip() for p in parts if p.strip()]

def canonicalize(items):
    # generic normalization: lowercase, strip punctuation, keep short tokens joined
    out = []
    for it in items:
        s = re.sub(r"[^a-zA-Z0-9\+\#\.\- ]", " ", it.lower()).strip()
        s = re.sub(r"\s+", " ", s)
        if s: out.append(s)
    # de-dupe preserve order
    seen, res = set(), []
    for x in out:
        if x not in seen:
            seen.add(x); res.append(x)
    return res

def extract_jd_facts(jd_text):
    prompt = f"""Extract concise JD facts as JSON.
- Return short, canonical skill names (e.g., 'java', 'spring boot', 'selenium', 'azure devops', 'oracle', 'pl/sql', 'spark').
- Ignore soft skills and responsibilities.
JD:
{jd_text}"""
    raw = call_json(prompt, JD_SCHEMA)
    return {
        "role_titles": canonicalize(to_list(raw.get("role_titles"))),
        "required_skills": canonicalize(to_list(raw.get("required_skills"))),
        "optional_skills": canonicalize(to_list(raw.get("optional_skills"))),
        "domain_keywords": canonicalize(to_list(raw.get("domain_keywords"))),
        "location_keywords": canonicalize(to_list(raw.get("location_keywords"))),
        "min_years": extract_min_years(jd_text),
    }

def extract_resume_facts(resume_text):
    prompt = f"""Extract concise resume facts as JSON.
- Return short, canonical skill names (e.g., 'java', 'oracle', 'selenium', 'jenkins').
- Ignore job responsibilities sentences.
Resume:
{resume_text}"""
    raw = call_json(prompt, RESUME_SCHEMA)
    return {
        "role_titles": canonicalize(to_list(raw.get("role_titles"))),
        "skills": canonicalize(to_list(raw.get("skills"))),
        "domain_keywords": canonicalize(to_list(raw.get("domain_keywords"))),
        "location_keywords": canonicalize(to_list(raw.get("location_keywords"))),
        "years_overall": extract_years_from_resume(resume_text),
    }

# ===================== Scoring =====================
WEIGHTS = {
    "required_overlap": 0.55,   # required skill hit
    "optional_overlap": 0.10,   # nice to have
    "title_align":      0.15,   # generic alignment
    "years_fit":        0.10,   # regex-based
    "cosine":           0.10,   # bag-of-words cosine
}

CALIB = {"scale": 1.02, "offset": 2}  # small bias to match typical GPT-ish scale

def jaccard(a, b):
    A, B = set(a), set(b)
    if not A and not B: return 0.0
    return len(A & B) / len(A | B)

def title_alignment(jd_titles, cv_titles):
    # token overlap of titles (generic)
    jd = set(tokenize(" ".join(jd_titles)))
    cv = set(tokenize(" ".join(cv_titles)))
    if not jd or not cv: return 0.5
    inter = len(jd & cv); uni = len(jd | cv)
    return inter / uni

def years_fit(jd_min, cv_years):
    if jd_min <= 0: return 1.0
    if cv_years <= 0: return 0.0
    return min(cv_years / jd_min, 1.0)

def compute_score(jd_text, resume_text, jd_facts, cv_facts):
    # Assemble skill sets
    jd_required = jd_facts["required_skills"]
    jd_optional = jd_facts["optional_skills"]
    cv_all      = cv_facts["skills"]

    req_overlap = jaccard(jd_required, cv_all)
    opt_overlap = jaccard(jd_optional, cv_all) if jd_optional else 0.0
    title_align = title_alignment(jd_facts["role_titles"], cv_facts["role_titles"])
    yrs_fit     = years_fit(jd_facts["min_years"], cv_facts["years_overall"])
    cos         = cosine_sim(bow(jd_text), bow(resume_text))

    raw = (
        WEIGHTS["required_overlap"] * req_overlap +
        WEIGHTS["optional_overlap"] * opt_overlap +
        WEIGHTS["title_align"]      * title_align +
        WEIGHTS["years_fit"]        * yrs_fit +
        WEIGHTS["cosine"]           * cos
    )
    score = raw * 100.0
    score = score * CALIB["scale"] + CALIB["offset"]
    score = int(max(0, min(100, round(score))))
    # what matched / missing (for explanation)
    matched = sorted(list(set(jd_required) & set(cv_all)))
    missing = sorted(list(set(jd_required) - set(cv_all)))
    return score, matched, missing

# ===================== Follow-ups =====================
def generate_followup(jd_text, resume_text, candidate_name):
    prompt = f"""
You are a recruiter at Terrabit Consulting writing to a candidate named {candidate_name}.
Return exactly three sections:

### WhatsApp Message to Candidate
<friendly outreach; mention inferred role; 2–3 slots; confirm phone/email if missing>

### Email to Candidate
Subject: Quick chat about a {{<role>}} opportunity at Terrabit Consulting
Dear {candidate_name},
<3–5 lines about fit based on their resume; ask for 2–3 preferred slots this week; polite close>

### Screening Questions (Tailored)
- <Q1 about JD core>
- <Q2 about a missing/weak skill>
- <Q3 about domain/tools in JD>
- <Q4 optional availability/notice>
- <Q5 optional location/remote>

Information to use:
JD:
{jd_text}

Resume:
{resume_text}
"""
    try:
        resp = gemini_text.generate_content(prompt)
        return (getattr(resp, "text", "") or "").strip()
    except Exception as e:
        return f"⚠️ Follow-up generation failed: {e}"

# ===================== Session State =====================
if "results" not in st.session_state: st.session_state["results"] = []
if "processed_resumes" not in st.session_state: st.session_state["processed_resumes"] = set()
if "jd_text" not in st.session_state: st.session_state["jd_text"] = ""
if "jd_file" not in st.session_state: st.session_state["jd_file"] = None
if "summary" not in st.session_state: st.session_state["summary"] = []

# Reset button → brand new session (so you can upload a completely different JD + resumes)
if st.button("🔁 Start New Matching Session"):
    st.session_state.clear()
    st.rerun()

# ===================== Uploaders =====================
jd_file = st.file_uploader("📄 Upload Job Description", type=["txt", "pdf", "docx"], key="jd_uploader")
resume_files = st.file_uploader("📑 Upload Candidate Resumes", type=["txt", "pdf", "docx"], accept_multiple_files=True, key="resume_uploader")

# Keep JD for this run
if jd_file and not st.session_state.get("jd_text"):
    st.session_state["jd_text"] = normalize_text(read_file(jd_file))
    st.session_state["jd_file"] = jd_file.name

jd_text = st.session_state.get("jd_text", "")

# ===================== Run Matching =====================
if st.button("🚀 Run Matching") and jd_text and resume_files:
    # extract JD facts once
    jd_facts = extract_jd_facts(jd_text)

    for resume_file in resume_files:
        if resume_file.name in st.session_state["processed_resumes"]:
            continue

        resume_text = normalize_text(read_file(resume_file))
        name = extract_candidate_name(resume_text, resume_file.name)
        email = extract_email(resume_text)

        with st.spinner(f"🔎 Analyzing {name}..."):
            cv_facts = extract_resume_facts(resume_text)
            score, matched, missing = compute_score(jd_text, resume_text, jd_facts, cv_facts)

        warn = "\n\n**Warning**: Score below 70% – candidate may not meet core requirements." if score < 70 else ""
        role_line = "Experience overlaps with JD focus" if title_alignment(jd_facts["role_titles"], cv_facts["role_titles"]) >= 0.5 else "Role emphasis differs from JD."
        gaps_line = []
        if score < 70:
            if jd_facts["min_years"] and cv_facts["years_overall"] < jd_facts["min_years"]:
                gaps_line.append(f"Insufficient years (JD: {jd_facts['min_years']}+, Resume: {cv_facts['years_overall']})")
            if missing:
                gaps_line.append(f"Missing required skills: {', '.join(missing)}")
        gaps_text = "; ".join(gaps_line) if gaps_line else ("—" if score >= 70 else "Missing core skills / insufficient years")

        result_md = f"""**Name**: {name}
**Score**: [{score}]%

**Reason**:
- **Role Match**: {role_line}
- **Skill Match**: Matched: {', '.join(matched) if matched else 'None'}. Missing: {', '.join(missing) if missing else 'None'}.
- **Major Gaps**: {gaps_text}{warn}
"""

        st.session_state["results"].append({
            "correct_name": name,
            "email": email,
            "score": score,
            "result": result_md,
            "resume_text": resume_text
        })
        st.session_state["processed_resumes"].add(resume_file.name)
        st.session_state["summary"].append({"Candidate Name": name, "Email": email, "Score": score})

# ===================== Results =====================
for entry in st.session_state["results"]:
    st.markdown("---")
    st.subheader(f"📌 {entry['correct_name']}")
    st.markdown(f"📧 **Email**: {entry['email']}")
    st.markdown(entry["result"], unsafe_allow_html=True)

    s = entry["score"]
    if s < 50:
        st.error("❌ Not suitable – Major role mismatch")
    elif s < 70:
        st.warning("⚠️ Consider with caution – Lacks core skills")
    else:
        st.success("✅ Strong match – Good alignment with JD")

    if st.button(f"✉️ Generate Follow-up for {entry['correct_name']}", key=f"followup_{entry['correct_name']}"):
        with st.spinner("Generating messages..."):
            followup = generate_followup(st.session_state["jd_text"], entry["resume_text"], entry["correct_name"])
        st.markdown("---")
        st.markdown(followup, unsafe_allow_html=True)

# ===================== Summary =====================
if st.session_state["summary"]:
    st.markdown("### 📊 Summary of All Candidates")
    df_summary = pd.DataFrame(st.session_state["summary"]).sort_values(by="Score", ascending=False)
    st.dataframe(df_summary)
    xbuf = io.BytesIO()
    with pd.ExcelWriter(xbuf, engine="openpyxl") as w:
        df_summary.to_excel(w, index=False)
    st.download_button("📥 Download Summary as Excel",
                       data=xbuf.getvalue(),
                       file_name="resume_match_summary.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ===================== Debug (optional) =====================
with st.expander("🔍 Debug: last JD/Resume facts"):
    if st.session_state.get("jd_text"):
        st.write("JD facts (min years auto):")
        try:
            st.json(extract_jd_facts(st.session_state["jd_text"]))
        except Exception:
            st.caption("Upload & Run to view.")
